import os
import io
from pathlib import Path
from urllib.parse import urlparse, urlunparse

import requests
from dotenv import load_dotenv
from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import StreamingResponse

load_dotenv()

router = APIRouter(prefix="/api", tags=["Gateway"])

ORCH_URL = os.getenv("ORCH_URL", "http://orchestrator_service:8001/run")
TASK_URL = os.getenv("TASK_URL", "http://task_service:8002/task")
ORCH_TIMEOUT = int(os.getenv("ORCH_TIMEOUT", "900"))
TASK_TIMEOUT = int(os.getenv("TASK_TIMEOUT", "20"))
ALLOWED_ATTACHMENT_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_ATTACHMENT_BYTES = int(os.getenv("MAX_ATTACHMENT_BYTES", str(8 * 1024 * 1024)))
MAX_ATTACHMENT_CHARS = int(os.getenv("MAX_ATTACHMENT_CHARS", "12000"))
MAX_TOTAL_ATTACHMENT_CHARS = int(os.getenv("MAX_TOTAL_ATTACHMENT_CHARS", "24000"))


def _with_path(service_url: str, path: str) -> str:
    parsed_url = urlparse(service_url)
    return urlunparse(parsed_url._replace(path=path, params="", query="", fragment=""))


def _orch_path(path: str) -> str:
    return _with_path(ORCH_URL, path)


def _task_path(path: str) -> str:
    return _with_path(TASK_URL, path)


def _forward_json(method: str, url: str, payload: dict | None = None, timeout: int = 30):
    try:
        response = requests.request(method, url, json=payload, timeout=timeout)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.ConnectionError as exc:
        raise HTTPException(status_code=503, detail=f"Cannot connect to service: {exc}") from exc
    except requests.exceptions.Timeout as exc:
        raise HTTPException(status_code=504, detail="Service request timed out") from exc
    except requests.exceptions.HTTPError as exc:
        status_code = exc.response.status_code if exc.response is not None else 502
        detail = exc.response.text if exc.response is not None else str(exc)
        raise HTTPException(status_code=status_code, detail=detail) from exc


def _truncate_text(text: str, max_chars: int) -> tuple[str, bool]:
    cleaned = "\n".join(line.rstrip() for line in str(text or "").splitlines()).strip()
    if len(cleaned) <= max_chars:
        return cleaned, False
    return cleaned[:max_chars].rstrip(), True


def _extract_txt(raw_bytes: bytes) -> str:
    try:
        return raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return raw_bytes.decode("latin-1", errors="replace")


def _extract_pdf(raw_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="PDF support is not installed") from exc

    try:
        reader = PdfReader(io.BytesIO(raw_bytes))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise HTTPException(status_code=400, detail="Encrypted PDF files are not supported") from exc
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read PDF attachment: {exc}") from exc


def _extract_docx(raw_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="DOCX support is not installed") from exc

    try:
        document = Document(io.BytesIO(raw_bytes))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
        return "\n".join(parts)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Could not read DOCX attachment: {exc}") from exc


def _extract_attachment(filename: str, content_type: str, raw_bytes: bytes, max_chars: int) -> dict:
    safe_name = Path(filename or "").name
    suffix = Path(safe_name).suffix.lower()

    if suffix not in ALLOWED_ATTACHMENT_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_ATTACHMENT_EXTENSIONS))
        raise HTTPException(status_code=400, detail=f"Only {allowed} attachments are supported")

    if len(raw_bytes) > MAX_ATTACHMENT_BYTES:
        max_mb = round(MAX_ATTACHMENT_BYTES / (1024 * 1024), 1)
        raise HTTPException(status_code=400, detail=f"Attachment is too large. Maximum size is {max_mb} MB")

    if suffix == ".txt":
        extracted_text = _extract_txt(raw_bytes)
    elif suffix == ".pdf":
        extracted_text = _extract_pdf(raw_bytes)
    else:
        extracted_text = _extract_docx(raw_bytes)

    text, truncated = _truncate_text(extracted_text, max_chars)
    if not text:
        raise HTTPException(status_code=400, detail=f"No readable text was found in {safe_name}")

    return {
        "filename": safe_name,
        "content_type": content_type,
        "size_bytes": len(raw_bytes),
        "text": text,
        "text_characters": len(text),
        "truncated": truncated,
    }


async def _read_execute_payload(request: Request) -> dict:
    content_type = request.headers.get("content-type", "").lower()
    if "multipart/form-data" not in content_type:
        payload = await request.json()
        goal = str(payload.get("goal", "")).strip()
        if not goal:
            raise HTTPException(status_code=400, detail="Goal is required")
        return {"goal": goal, "attachments": payload.get("attachments") or []}

    form = await request.form()
    goal = str(form.get("goal") or "").strip()
    files = []
    for key in ("attachments", "attachment", "file"):
        files.extend(item for item in form.getlist(key) if getattr(item, "filename", ""))

    if not goal and files:
        goal = "Create a software development plan from the attached document."
    if not goal:
        raise HTTPException(status_code=400, detail="Goal is required")

    attachments = []
    remaining_chars = MAX_TOTAL_ATTACHMENT_CHARS
    for upload in files:
        raw_bytes = await upload.read()
        attachment = _extract_attachment(
            upload.filename,
            getattr(upload, "content_type", "") or "",
            raw_bytes,
            max(0, min(MAX_ATTACHMENT_CHARS, remaining_chars)),
        )
        attachments.append(attachment)
        remaining_chars -= attachment["text_characters"]
        if remaining_chars <= 0:
            break

    return {"goal": goal, "attachments": attachments}


@router.post("/execute")
async def execute_goal(request: Request):
    payload = await _read_execute_payload(request)
    return _forward_json("POST", ORCH_URL, payload, ORCH_TIMEOUT)


@router.post("/execute/stream")
async def execute_goal_stream(request: Request):
    payload = await _read_execute_payload(request)

    def event_generator():
        try:
            with requests.post(
                _orch_path("/run/stream"),
                json=payload,
                timeout=ORCH_TIMEOUT,
                stream=True,
            ) as response:
                if response.status_code != 200:
                    yield (
                        "event: error\n"
                        f"data: {{\"error\": \"Orchestrator stream failed with {response.status_code}\"}}\n\n"
                    )
                    return

                for line in response.iter_lines(decode_unicode=True):
                    if line is None:
                        continue
                    yield f"{line}\n"
        except GeneratorExit:
            print("Client disconnected from gateway stream")
            raise
        except requests.exceptions.ConnectionError as exc:
            yield f"event: error\ndata: {{\"error\": \"Cannot connect to orchestrator: {exc}\"}}\n\n"
        except requests.exceptions.Timeout:
            yield f"event: error\ndata: {{\"error\": \"Orchestrator timed out after {ORCH_TIMEOUT}s\"}}\n\n"
        except Exception as exc:
            yield f"event: error\ndata: {{\"error\": \"Streaming failed: {exc}\"}}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.get("/kpis")
def get_kpis():
    return _forward_json("GET", _task_path("/kpis"), timeout=TASK_TIMEOUT)


@router.get("/tasks/{task_id}")
def get_task(task_id: str):
    return _forward_json("GET", _task_path(f"/tasks/{task_id}"), timeout=TASK_TIMEOUT)


@router.get("/tasks/{task_id}/events")
def get_task_events(task_id: str):
    return _forward_json("GET", _task_path(f"/tasks/{task_id}/events"), timeout=TASK_TIMEOUT)


@router.post("/tasks/{task_id}/approve-plan")
def approve_plan(task_id: str, payload: dict | None = None):
    return _forward_json(
        "POST",
        _orch_path(f"/tasks/{task_id}/approve-plan"),
        payload or {},
        ORCH_TIMEOUT,
    )


@router.post("/tasks/{task_id}/edit-plan")
def edit_plan(task_id: str, payload: dict):
    return _forward_json("POST", _orch_path(f"/tasks/{task_id}/edit-plan"), payload, ORCH_TIMEOUT)


@router.post("/tasks/{task_id}/request-revision")
def request_revision(task_id: str, payload: dict | None = None):
    return _forward_json(
        "POST",
        _orch_path(f"/tasks/{task_id}/request-revision"),
        payload or {},
        ORCH_TIMEOUT,
    )


@router.post("/tasks/{task_id}/approve-output")
def approve_output(task_id: str, payload: dict | None = None):
    return _forward_json(
        "POST",
        _orch_path(f"/tasks/{task_id}/approve-output"),
        payload or {},
        ORCH_TIMEOUT,
    )


@router.post("/tasks/{task_id}/reject-output")
def reject_output(task_id: str, payload: dict | None = None):
    return _forward_json(
        "POST",
        _orch_path(f"/tasks/{task_id}/reject-output"),
        payload or {},
        ORCH_TIMEOUT,
    )


@router.post("/tasks/{task_id}/complete")
def complete_task(task_id: str, payload: dict | None = None):
    return _forward_json(
        "POST",
        _orch_path(f"/tasks/{task_id}/complete"),
        payload or {},
        ORCH_TIMEOUT,
    )


@router.get("/health")
def gateway_health():
    try:
        orch_response = requests.get(_orch_path("/health"), timeout=10)
        task_response = requests.get(_task_path("/health"), timeout=10)

        return {
            "status": "healthy",
            "orchestrator": orch_response.json()
            if orch_response.status_code == 200
            else "unhealthy",
            "task_service": task_response.json()
            if task_response.status_code == 200
            else "unhealthy",
        }
    except Exception as exc:
        return {"status": "unhealthy", "error": str(exc)}
